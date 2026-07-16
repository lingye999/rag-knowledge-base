import re
import os
import io
import numpy as np
import jieba
import docx
import fitz
import pdfplumber


def read_file(path: str, force_ocr: bool = False) -> str:
    """根据文件后缀自动选择读取方式，支持 txt/docx/pdf"""
    if not os.path.exists(path):
        raise FileNotFoundError(f"文件不存在: {path}")

    if path.endswith(".txt"):
        for encoding in ["utf-8", "gbk", "gb2312", "gb18030"]:
            try:
                with open(path, "r", encoding=encoding) as f:
                    return f.read()
            except (UnicodeDecodeError, UnicodeError):
                continue
        with open(path, "r", encoding="utf-8") as f:
            return f.read()

    elif path.endswith(".docx"):
        return read_docx(path)
    elif path.endswith(".pdf"):
        return read_pdf(path, force_ocr=force_ocr)
    else:
        raise ValueError(f"不支持的文件格式: {path}，仅支持 .txt / .docx / .pdf")


def read_docx(path: str) -> str:
    """读取word文本，包括段落和表格"""
    try:
        doc = docx.Document(path)
    except Exception as e:
        raise RuntimeError(f"无法读取 DOCX 文件（可能已损坏、加密或格式为旧版 .doc）: {path}") from e

    paragraph = []
    for para in doc.paragraphs:
        if para.text.strip():
            paragraph.append(para.text)

    # 同时提取表格中的文字
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraph.append(cell.text)

    result = "\n".join(paragraph)
    if not result.strip():
        print(f"[警告] DOCX 文件中未提取到文字（可能全是图片）: {path}")
    return result


def read_pdf(path: str, force_ocr: bool = False) -> str:
    """读取pdf文本：默认 pdfplumber 优先，force_ocr=True 时优先 OCR"""
    if force_ocr:
        # ── OCR 优先模式 ──────────────────────────────────
        result = _read_pdf_ocr(path)
        if result and result.strip():
            return result
        # OCR 失败，退回 pdfplumber
        result = _read_pdf_plumber(path)
        if result and result.strip():
            return result
        raise RuntimeError(f"无法提取 PDF 文字（已尝试 OCR + 文本提取）: {path}")

    # ── 默认模式：pdfplumber 优先 ────────────────────────
    result = _read_pdf_plumber(path)
    if result and result.strip():
        return result

    # OCR 兜底
    result = _read_pdf_ocr(path)
    if result and result.strip():
        return result

    raise RuntimeError(f"无法提取 PDF 文字（已尝试文本提取 + OCR）: {path}")


def _read_pdf_plumber(path: str) -> str:
    """pdfplumber 文本提取"""
    try:
        with pdfplumber.open(path) as pdf:
            pages = []
            empty_pages = 0
            total_pages = len(pdf.pages)
            for page in pdf.pages:
                text = page.extract_text()
                page_text = text.strip() if text else ""
                is_garbage = _is_garbage_page(page_text)

                if not is_garbage:
                    if page_text:
                        pages.append(page_text)
                else:
                    empty_pages += 1

                if not is_garbage:
                    tables = page.extract_tables()
                    for table in tables:
                        rows = []
                        for row in table:
                            cells = [c.strip() if c else "" for c in row]
                            line = " | ".join(cells)
                            noise_ratio = sum(1 for c in line if c in "√/|0123456789 ") / len(line) if line else 1
                            if noise_ratio > 0.6:
                                continue
                            rows.append(line)
                        if rows:
                            pages.append("\n".join(rows))

            result = "\n".join(pages)
            result = clean_table_noise(result)
            empty_ratio = empty_pages / total_pages if total_pages > 0 else 1
            if result.strip() and empty_ratio < 0.5:
                return result
    except Exception:
        pass
    return ""


def _read_pdf_ocr(path: str) -> str:
    """EasyOCR 图片识别"""
    try:
        import easyocr

        doc = fitz.open(path)
        reader = easyocr.Reader(['ch_sim', 'en'])
        pages = []
        for page in doc:
            pix = page.get_pixmap(dpi=200)
            img = np.frombuffer(pix.samples, dtype=np.uint8).reshape(pix.height, pix.width, pix.n)
            result = reader.readtext(img)
            line_texts = [text for _, text, conf in result if conf > 0.3]
            pages.append("".join(line_texts))

        doc.close()
        result = "\n".join(pages)
        result = clean_table_noise(result)
        return result if result.strip() else ""
    except Exception:
        return ""


def chunk_by_sentence(text: str) -> list[str]:
    """按照【。！？.!?】分割句子，支持中英文"""

    sentences = re.split(r"[。！？.!?]", text)
    return [s.strip() for s in sentences if s.strip()]


def chunk_by_paragraph(text: str) -> list[str]:
    """按照段落去分割这个文本"""

    paragraph = text.split("\n")
    return [s.strip() for s in paragraph if s.strip()]


def chunk_by_size(text: str, chunk_size: int = 200, overlap: int = 50) -> list[str]:

    chunks = []  # 存储这个截断的语句
    start = 0

    while start < len(text):
        end = start + chunk_size
        chunk = text[start:end]
        chunks.append(chunk)
        start += chunk_size - overlap  # 减去重叠的部分
    return chunks


def chunk_by_jieba(text: str, max_words: int = 50) -> list[str]:
    """按中文结果分词分块"""
    words = jieba.lcut(text)  # 分词
    chunks = []
    for i in range(0, len(words), max_words):
        chunk = "".join(words[i:i + max_words])
        chunks.append(chunk)
    return chunks


def chunk_text(text: str, method: str = "auto") -> list[str]:
    """统一分块调度，支持 auto 自动选择"""
    # 空文本直接返回空列表
    if not text or not text.strip():
        print("[警告] 文本为空，无内容可分块")
        return []

    if method == "auto":
        # 1. 数段落
        paragraphs = [p.strip() for p in text.split("\n") if p.strip()]

        # 2. 算中文占比
        total_chars = len(text.strip())
        if total_chars > 0:
            chinese_chars = len(re.findall(r'[一-鿿]', text))
            chinese_ratio = chinese_chars / total_chars
        else:
            chinese_ratio = 0

        # 3. 自动选择
        if chinese_ratio > 0.3:
            method = "jieba"
        elif len(paragraphs) >= 3:
            method = "paragraph"
        else:
            method = "sentence"

        print(f"[Auto] 中文占比 {chinese_ratio:.0%}，段落 {len(paragraphs)} 个，使用 {method} 分块")

    # 4. 调度
    if method == "sentence":
        return chunk_by_sentence(text)
    elif method == "paragraph":
        return chunk_by_paragraph(text)
    elif method == "jieba":
        return chunk_by_jieba(text)
    elif method == "size":
        return chunk_by_size(text)
    else:
        raise ValueError(f"不支持的分块方法: {method}，可选: sentence / paragraph / jieba / size / auto")


def clean_table_noise(text: str) -> str:
    """清洗 pdfplumber 表格提取产生的脏数据"""
    import re
    lines = text.split("\n")
    cleaned = []
    for line in lines:
        s = line.strip()
        # 去掉空行
        if not s:
            continue
        # 去掉纯符号行（√ / | 等）
        if all(c in "√/| " for c in s):
            continue
        # 去掉" | " 和 "/ " 这种单元格分隔符，保留文字
        s = s.replace(" | ", " ").replace("/ ", " ")
        # 去掉孤立的 "|"
        s = s.replace(" |", "").replace("| ", "")
        # 去掉行首行尾的 |（表格边框残余）
        s = s.strip("| ")
        # 去掉超短的碎片行（≤3个字符且不含中文）
        if len(s) <= 3 and not re.search(r'[\u4e00-\u9fff]', s):
            continue
        # 去掉一半以上都是数字+符号的行（技术图纸标注）
        no_alpha = sum(1 for c in s if c in "0123456789/\\-CDVYEABPJ ")
        if len(s) > 3 and no_alpha / len(s) > 0.7:
            continue
        if s:
            cleaned.append(s)
    return "\n".join(cleaned)


def _is_garbage_page(text: str) -> bool:
    """判断一页提取的文本是否是垃圾（技术图纸标注），是则跳过"""
    if not text or not text.strip():
        return True
    # 去除空白后长度
    stripped = text.strip()
    if len(stripped) < 20:
        return False  # 短文本交给 empty_ratio 判断

    # 中文占比 < 15% → 技术图纸
    chinese = sum(1 for c in stripped if '\u4e00' <= c <= '\u9fff')
    ch_ratio = chinese / len(stripped)
    if ch_ratio < 0.15:
        return True

    # "单词"平均字符数 < 2（全是碎片，如 "流 aking cu 2 rr 0 en"）
    words = stripped.split()
    if words:
        avg_word_len = sum(len(w) for w in words) / len(words)
        if avg_word_len < 2.5:
            return True

    return False


def filter_stopwords(words: list[str]) -> list[str]:
    """去掉常见的停用词（的、了、是、在...）"""
    stopwords = {"的", "了", "是", "在", "和", "就", "都", "而", "及", "与",
                 "着", "或", "一个", "没有", "我们", "你们", "他们", "它",
                 "她", "他", "有", "不", "被", "把", "这", "那", "也"}
    return [w for w in words if w not in stopwords and w.strip()]

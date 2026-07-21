"""文档读取：txt / docx / pdf 多格式支持"""
import os
import docx
from .plumber import read_pdf_plumber
from .ocr import read_pdf_ocr
from .marker_reader import read_pdf_marker
from .hybrid_reader import read_pdf_hybrid


def read_file(path: str, force_ocr: bool = False, use_marker: bool = False) -> str:
    """根据文件后缀自动选择读取方式，支持 txt/docx/pdf

    Args:
        use_marker: 复杂 PDF 手动启用 Marker（需要 GPU + 本地模型）
    """
    if not os.path.exists(path):
        raise FileNotFoundError(f"文件不存在: {path}")

    if path.endswith(".txt"):
        for encoding in ["utf-8", "gbk", "gb2312", "gb18030"]:
            try:
                with open(path, "r", encoding=encoding) as f:
                    return f.read()
            except (UnicodeDecodeError, UnicodeError):
                continue
        with open(path, "r", encoding="utf-8") as f:
            return f.read()

    elif path.endswith(".docx"):
        return _read_docx(path)
    elif path.endswith(".pdf"):
        return _read_pdf(path, force_ocr=force_ocr, use_marker=use_marker)
    else:
        raise ValueError(f"不支持的文件格式: {path}，仅支持 .txt / .docx / .pdf")


def _read_docx(path: str) -> str:
    """读取word文本，包括段落和表格"""
    try:
        doc = docx.Document(path)
    except Exception as e:
        raise RuntimeError(f"无法读取 DOCX 文件: {path}") from e

    paragraph = []
    for para in doc.paragraphs:
        if para.text.strip():
            paragraph.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraph.append(cell.text)

    result = "\n".join(paragraph)
    if not result.strip():
        print(f"[警告] DOCX 文件中未提取到文字: {path}")
    return result


def _read_pdf(path: str, force_ocr: bool = False, use_marker: bool = False) -> str:
    """读取 PDF 文本

    优先级：
      1. Marker（仅 use_marker=True 时，用于复杂 PDF）
      2. 轻量混合（默认：逐行判断乱码 → 选择性 OCR）
      3. 纯 OCR 兜底
    """
    # ── Marker 模式（用户手动启用，用于复杂排版/表格 PDF）──
    if use_marker:
        try:
            result = read_pdf_marker(path)
            if result and result.strip():
                print(f"[PDF] Marker 解析成功: {path}")
                return result
        except Exception as e:
            print(f"[PDF] Marker 解析失败，降级: {e}")

    # ── OCR 强制模式 ──
    if force_ocr:
        result = read_pdf_ocr(path)
        if result and result.strip():
            return result
        result = read_pdf_plumber(path)
        if result and result.strip():
            return result
        raise RuntimeError(f"无法提取 PDF 文字（已尝试 OCR + 文本提取）: {path}")

    # ── 默认：轻量混合（逐块乱码检测 + OCR 替换）──
    try:
        result = read_pdf_hybrid(path)
        if result and result.strip():
            return result
    except Exception as e:
        print(f"[PDF] 混合解析失败，降级: {e}")

    # ── 兜底：纯 OCR ──
    result = read_pdf_ocr(path)
    if result and result.strip():
        return result

    raise RuntimeError(f"无法提取 PDF 文字（已尝试混合解析 → OCR）: {path}")

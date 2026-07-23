"""文档读取：txt / docx / pdf 多格式支持"""
import os
import docx
from .plumber import read_pdf_plumber
from .ocr import read_pdf_ocr, read_pdf_ocr_pages
from .marker_reader import read_pdf_marker
from .hybrid_reader import read_pdf_hybrid, read_pdf_hybrid_pages
from .parse_result import ParseResult, TextBlock


def read_file(path: str, force_ocr: bool = False, use_marker: bool = False) -> str:
    """根据文件后缀自动选择读取方式，支持 txt/docx/pdf

    Args:
        use_marker: 复杂 PDF 手动启用 Marker（需要 GPU + 本地模型）
    """
    return read_file_structured(
        path,
        force_ocr=force_ocr,
        use_marker=use_marker,
    ).text


def read_file_structured(
    path: str,
    force_ocr: bool = False,
    use_marker: bool = False,
    page_numbers: set[int] | None = None,
    allow_ocr: bool = True,
) -> ParseResult:
    """读取文件并返回结构化解析结果，兼容后续 page/bbox/confidence 扩展。"""
    if not os.path.exists(path):
        raise FileNotFoundError(f"文件不存在: {path}")

    if path.endswith(".txt"):
        text = ""
        for encoding in ["utf-8", "gbk", "gb2312", "gb18030"]:
            try:
                with open(path, "r", encoding=encoding) as f:
                    text = f.read()
                    break
            except (UnicodeDecodeError, UnicodeError):
                continue
        if not text:
            with open(path, "r", encoding="utf-8") as f:
                text = f.read()
        return _single_block_result(path, text, parser="txt", source="txt")

    elif path.endswith(".docx"):
        text = _read_docx(path)
        return _single_block_result(path, text, parser="docx", source="docx")
    elif path.endswith(".pdf"):
        return _read_pdf_structured(
            path,
            force_ocr=force_ocr,
            use_marker=use_marker,
            page_numbers=page_numbers,
            allow_ocr=allow_ocr,
        )
    else:
        raise ValueError(f"不支持的文件格式: {path}，仅支持 .txt / .docx / .pdf")


def _single_block_result(
    path: str,
    text: str,
    parser: str,
    source: str,
) -> ParseResult:
    blocks = []
    if text and text.strip():
        blocks.append(TextBlock(text=text, source=source))
    return ParseResult(path=path, blocks=blocks, parser=parser)


def _read_docx(path: str) -> str:
    """读取word文本，包括段落和表格"""
    try:
        doc = docx.Document(path)
    except Exception as e:
        raise RuntimeError(f"无法读取 DOCX 文件: {path}") from e

    paragraph = []
    for para in doc.paragraphs:
        if para.text.strip():
            paragraph.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraph.append(cell.text)

    result = "\n".join(paragraph)
    if not result.strip():
        print(f"[警告] DOCX 文件中未提取到文字: {path}")
    return result


def _read_pdf_structured(
    path: str,
    force_ocr: bool = False,
    use_marker: bool = False,
    page_numbers: set[int] | None = None,
    allow_ocr: bool = True,
) -> ParseResult:
    if use_marker:
        text = _read_pdf(path, force_ocr=False, use_marker=True)
        return _single_block_result(
            path, text, parser="pdf_marker", source="pdf_marker"
        )

    if force_ocr:
        pages = read_pdf_ocr_pages(path, page_numbers=page_numbers)
        if pages and any(page.strip() for page in pages):
            return _page_blocks_result(path, pages, parser="pdf_ocr")
        text = read_pdf_plumber(path)
        if text.strip():
            return _single_block_result(
                path, text, parser="pdf_plumber_fallback", source="pdf_plumber"
            )
        raise RuntimeError(f"Unable to extract PDF text: {path}")

    try:
        pages = read_pdf_hybrid_pages(
            path, page_numbers=page_numbers, allow_ocr=allow_ocr
        )
        if pages and any(page.strip() for page in pages):
            return _page_blocks_result(path, pages, parser="pdf_hybrid")
    except Exception as exc:
        print(f"[PDF] Hybrid parse failed, falling back to OCR: {exc}")

    if not allow_ocr:
        raise RuntimeError(f"Hybrid parse produced no text without OCR: {path}")
    pages = read_pdf_ocr_pages(path, page_numbers=page_numbers)
    if pages and any(page.strip() for page in pages):
        return _page_blocks_result(path, pages, parser="pdf_ocr_fallback")
    raise RuntimeError(f"Unable to extract PDF text: {path}")


def _page_blocks_result(path: str, pages: list[str], parser: str) -> ParseResult:
    blocks = [
        TextBlock(text=text, source=parser, page=page_number)
        for page_number, text in enumerate(pages, start=1)
        if text.strip()
    ]
    return ParseResult(path=path, blocks=blocks, parser=parser)


def _read_pdf(path: str, force_ocr: bool = False, use_marker: bool = False) -> str:
    """读取 PDF 文本

    优先级：
      1. Marker（仅 use_marker=True 时，用于复杂 PDF）
      2. 轻量混合（默认：逐行判断乱码 → 选择性 OCR）
      3. 纯 OCR 兜底
    """
    # ── Marker 模式（用户手动启用，用于复杂排版/表格 PDF）──
    if use_marker:
        try:
            result = read_pdf_marker(path)
            if result and result.strip():
                print(f"[PDF] Marker 解析成功: {path}")
                return result
        except Exception as e:
            print(f"[PDF] Marker 解析失败，降级: {e}")

    # ── OCR 强制模式 ──
    if force_ocr:
        result = read_pdf_ocr(path)
        if result and result.strip():
            return result
        result = read_pdf_plumber(path)
        if result and result.strip():
            return result
        raise RuntimeError(f"无法提取 PDF 文字（已尝试 OCR + 文本提取）: {path}")

    # ── 默认：轻量混合（逐块乱码检测 + OCR 替换）──
    try:
        result = read_pdf_hybrid(path)
        if result and result.strip():
            return result
    except Exception as e:
        print(f"[PDF] 混合解析失败，降级: {e}")

    # ── 兜底：纯 OCR ──
    result = read_pdf_ocr(path)
    if result and result.strip():
        return result

    raise RuntimeError(f"无法提取 PDF 文字（已尝试混合解析 → OCR）: {path}")
